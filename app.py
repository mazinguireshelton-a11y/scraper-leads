"""
MOTOR DE BUSCA B2B UNIVERSAL COM IA (OpenRouter) - EDICAO PREMIUM
-----------------------------------------------------------------
"""

import streamlit as st
import pandas as pd
import requests
import re
import os
from io import BytesIO
from bs4 import BeautifulSoup
from datetime import date
from docx import Document
from reportlab.lib.pagesizes import letter
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
from reportlab.lib.styles import getSampleStyleSheet
from supabase import create_client, ClientOptions

# 1. Configuração da página (Primeiro comando Streamlit)
st.set_page_config(
    page_title="Mira", 
    page_icon="⚡", 
    layout="wide",
    initial_sidebar_state="collapsed"
)

# Tenta carregar dotenv em ambiente local
try:
    from dotenv import load_dotenv
    load_dotenv()
except ImportError:
    pass

# --- COOKIES: mantém a sessão ligada mesmo depois de recarregar a página ---
from streamlit_cookies_manager import EncryptedCookieManager

try:
    _cookie_password = st.secrets.get("COOKIES_PASSWORD", os.getenv("COOKIES_PASSWORD", "mira-troque-esta-chave-2026"))
except Exception:
    _cookie_password = os.getenv("COOKIES_PASSWORD", "mira-troque-esta-chave-2026")

cookies = EncryptedCookieManager(prefix="mira_", password=_cookie_password)
if not cookies.ready():
    st.stop()

# --- PONTE JS: o login com Google (OAuth) devolve o token depois do "#" na URL,
# que o Python não consegue ler. Este script converte isso em query params (?),
# que o Streamlit já lê. Só entra em ação quando existe um "access_token" no link.
import streamlit.components.v1 as components
components.html("""
<script>
try {
    var loc = window.top.location;
    if (loc.hash && loc.hash.includes('access_token')) {
        var params = new URLSearchParams(loc.hash.substring(1));
        var newUrl = loc.pathname + '?' + params.toString();
        window.top.location.href = newUrl;
    }
} catch (e) {
    console.error('Mira OAuth bridge error:', e);
}
</script>
""", height=0)

# A recuperação de senha usa token_hash como query param normal (?token_hash=...&type=recovery),
# configurado no template de e-mail do Supabase. Isso evita problemas de segurança do
# navegador que impediam a leitura do link antigo (baseado em #access_token).

# --- CONFIGURAÇÃO DE CHAVES DE API ---
try:
    RAPIDAPI_KEY = st.secrets.get("RAPIDAPI_KEY", os.getenv("RAPIDAPI_KEY", ""))
    OPENROUTER_API_KEY = st.secrets.get("OPENROUTER_API_KEY", os.getenv("OPENROUTER_API_KEY", ""))
    GOOGLE_PLACES_API_KEY = st.secrets.get("GOOGLE_PLACES_API_KEY", os.getenv("GOOGLE_PLACES_API_KEY", ""))
    SUPABASE_URL = st.secrets.get("SUPABASE_URL", os.getenv("SUPABASE_URL", ""))
    SUPABASE_SERVICE_KEY = st.secrets.get("SUPABASE_SERVICE_KEY", os.getenv("SUPABASE_SERVICE_KEY", ""))
except Exception:
    RAPIDAPI_KEY = os.getenv("RAPIDAPI_KEY", "")
    OPENROUTER_API_KEY = os.getenv("OPENROUTER_API_KEY", "")
    GOOGLE_PLACES_API_KEY = os.getenv("GOOGLE_PLACES_API_KEY", "")
    SUPABASE_URL = os.getenv("SUPABASE_URL", "")
    SUPABASE_SERVICE_KEY = os.getenv("SUPABASE_SERVICE_KEY", "")

SEARCH_URL = "https://local-business-data.p.rapidapi.com/search"
OPENROUTER_API_URL = "https://openrouter.ai/api/v1/chat/completions"
LIMITE_DIARIO_PADRAO = 3
MAX_LEADS_FREE = 6
EMAILS_ADMIN = ["mazinguireshelton@gmail.com"]

# --- CSS PERSONALIZADO E ANIMAÇÕES LEVES ---
st.markdown("""
<style>
@import url('https://fonts.googleapis.com/css2?family=Plus+Jakarta+Sans:wght@400;500;600;700&family=JetBrains+Mono:wght@400;500&display=swap');

/* Reset e Variáveis Globais */
:root {
    --bg-dark: #090D16;
    --card-bg: rgba(22, 27, 38, 0.7);
    --border-color: rgba(255, 255, 255, 0.08);
    --accent: #3B82F6;
    --accent-glow: rgba(59, 130, 246, 0.35);
    --text-primary: #F3F4F6;
    --text-secondary: #9CA3AF;
}

/* Animações CSS Puras */
@keyframes fadeInUp {
    from {
        opacity: 0;
        transform: translateY(12px);
    }
    to {
        opacity: 1;
        transform: translateY(0);
    }
}

@keyframes pulseGlow {
    0% { box-shadow: 0 0 5px var(--accent-glow); }
    50% { box-shadow: 0 0 18px var(--accent-glow); }
    100% { box-shadow: 0 0 5px var(--accent-glow); }
}

.stApp {
    background-color: var(--bg-dark);
    color: var(--text-primary);
    font-family: 'Plus Jakarta Sans', sans-serif;
}

/* Aplicação de Animação em Blocos Principais */
div[data-testid="stForm"], 
div[data-testid="stExpander"], 
.stDataFrame, 
.stButton>button {
    animation: fadeInUp 0.4s cubic-bezier(0.16, 1, 0.3, 1) forwards;
}

/* Cabeçalhos */
h1, h2, h3 {
    font-family: 'Plus Jakarta Sans', sans-serif !important;
    font-weight: 700 !important;
    letter-spacing: -0.02em !important;
    color: #FFFFFF !important;
}

/* Estilo Glassmorphism nos Inputs */
.stTextInput input, .stTextArea textarea, .stNumberInput input {
    background: rgba(15, 20, 30, 0.6) !important;
    border: 1px solid var(--border-color) !important;
    border-radius: 10px !important;
    color: var(--text-primary) !important;
    font-family: 'Plus Jakarta Sans', sans-serif !important;
    backdrop-filter: blur(8px);
    transition: all 0.25s ease !important;
}

.stTextInput input:focus, .stTextArea textarea:focus {
    border-color: var(--accent) !important;
    box-shadow: 0 0 12px var(--accent-glow) !important;
}

/* Botão Principal com Efeito Iluminado */
.stButton>button[kind="primary"] {
    background: linear-gradient(135deg, #2563EB 0%, #1D4ED8 100%) !important;
    color: #FFFFFF !important;
    border: none !important;
    border-radius: 10px !important;
    font-weight: 600 !important;
    padding: 0.6rem 1.2rem !important;
    transition: all 0.2s ease !important;
    animation: pulseGlow 3s infinite;
}

.stButton>button[kind="primary"]:hover {
    transform: translateY(-2px);
    box-shadow: 0 6px 20px var(--accent-glow) !important;
}

/* Botões Secundários */
.stButton>button {
    background: var(--card-bg) !important;
    border: 1px solid var(--border-color) !important;
    border-radius: 10px !important;
    color: var(--text-primary) !important;
    backdrop-filter: blur(10px);
    transition: all 0.2s ease !important;
}

.stButton>button:hover {
    border-color: var(--accent) !important;
    color: var(--accent) !important;
}

/* Expanders (Cards de Propostas) com Vidro Fosco */
div[data-testid="stExpander"] {
    background: var(--card-bg) !important;
    border: 1px solid var(--border-color) !important;
    border-radius: 12px !important;
    backdrop-filter: blur(12px);
    margin-bottom: 0.8rem;
}

/* Barra Lateral */
section[data-testid="stSidebar"] {
    background-color: rgba(12, 16, 25, 0.95) !important;
    border-right: 1px solid var(--border-color) !important;
}

/* Estilização para Dispositivos Móveis */
@media (max-width: 768px) {
    .stApp { padding: 0.5rem; }
    .stButton>button { width: 100%; }
}
</style>
""", unsafe_allow_html=True)

# --- BANCO DE DADOS SUPABASE ---
@st.cache_resource
def get_supabase():
    if not SUPABASE_URL or not SUPABASE_SERVICE_KEY:
        return None
    try:
        opts = ClientOptions(flow_type="implicit")
        return create_client(SUPABASE_URL, SUPABASE_SERVICE_KEY, options=opts)
    except Exception:
        # Se a versão da lib não aceitar esse parâmetro, cai pro padrão
        return create_client(SUPABASE_URL, SUPABASE_SERVICE_KEY)

def garantir_registro_cliente(user_id, email, apelido=None, empresa=None, perfil_oferta=None):
    sb = get_supabase()
    if not sb: return
    existente = sb.table("clientes").select("*").eq("id", user_id).execute()
    if not existente.data:
        sb.table("clientes").insert({
            "id": user_id, "email": email, "apelido": apelido or email.split("@")[0],
            "empresa": empresa or "", "perfil_oferta": perfil_oferta or "",
            "limite_diario": LIMITE_DIARIO_PADRAO
        }).execute()

def buscar_apelido(user_id, email):
    sb = get_supabase()
    if not sb: return email.split("@")[0]
    resp = sb.table("clientes").select("apelido").eq("id", user_id).execute()
    if resp.data and resp.data[0].get("apelido"):
        return resp.data[0]["apelido"]
    return email.split("@")[0]

def buscar_perfil_completo(user_id):
    sb = get_supabase()
    if not sb: return {"apelido": "", "empresa": "", "perfil_oferta": ""}
    resp = sb.table("clientes").select("apelido, empresa, perfil_oferta").eq("id", user_id).execute()
    if resp.data:
        d = resp.data[0]
        return {"apelido": d.get("apelido") or "", "empresa": d.get("empresa") or "", "perfil_oferta": d.get("perfil_oferta") or ""}
    return {"apelido": "", "empresa": "", "perfil_oferta": ""}

def salvar_perfil_completo(user_id, apelido, empresa, perfil_oferta):
    sb = get_supabase()
    if not sb: return
    sb.table("clientes").update({"apelido": apelido, "empresa": empresa, "perfil_oferta": perfil_oferta}).eq("id", user_id).execute()

def verificar_e_registrar_uso(user_id, email=""):
    if email in EMAILS_ADMIN:
        return True, "Acesso Administrador ativado."

    sb = get_supabase()
    if not sb:
        return True, "Modo de demonstração sem limites ativado."

    cliente = sb.table("clientes").select("limite_diario, limite_total, plano").eq("id", user_id).execute()

    # Premium: sem limite nenhum (nem busca, nem quantidade, nem IA)
    plano = cliente.data[0].get("plano", "free") if cliente.data else "free"
    if plano == "premium":
        return True, "✨ Plano Premium — sem limites."

    limite_diario = cliente.data[0]["limite_diario"] if cliente.data else LIMITE_DIARIO_PADRAO
    limite_total = cliente.data[0].get("limite_total") if cliente.data else None

    # Limite total fixo (ex: beta testers com 3 tentativas, nunca reseta)
    if limite_total is not None:
        todos_registros = sb.table("uso_diario").select("contagem").eq("user_id", user_id).execute()
        usado_total = sum(r["contagem"] for r in (todos_registros.data or []))
        if usado_total >= limite_total:
            return False, f"Atingiste o limite de {limite_total} tentativas do modo beta. Obrigado por testar! Fala connosco para continuar a usar."

    hoje = str(date.today())
    registro = sb.table("uso_diario").select("*").eq("user_id", user_id).eq("data", hoje).execute()

    if registro.data:
        contagem = registro.data[0]["contagem"]
        if limite_total is None and contagem >= limite_diario:
            return False, f"Limite diário de {limite_diario} buscas atingido. Upgrade para Premium para buscas ilimitadas."
        sb.table("uso_diario").update({"contagem": contagem + 1}).eq("user_id", user_id).eq("data", hoje).execute()
    else:
        sb.table("uso_diario").insert({"user_id": user_id, "data": hoje, "contagem": 1}).execute()

    if limite_total is not None:
        restantes = limite_total - (usado_total + 1)
        return True, f"{restantes} tentativa(s) beta restante(s)."
    else:
        restantes = limite_diario - (registro.data[0]["contagem"] + 1 if registro.data else 1)
        return True, f"{restantes} buscas restantes hoje."

def buscar_plano_cliente(user_id):
    sb = get_supabase()
    if not sb: return "free"
    resp = sb.table("clientes").select("plano").eq("id", user_id).execute()
    return resp.data[0].get("plano", "free") if resp.data else "free"

def marcar_ia_usada(user_id):
    sb = get_supabase()
    if not sb: return
    sb.table("clientes").update({"ia_usado": True}).eq("id", user_id).execute()

def verificar_ia_disponivel(user_id, email=""):
    """Retorna (permitido: bool, mensagem: str). No free, a IA só pode ser usada 1 vez para sempre."""
    if email in EMAILS_ADMIN:
        return True, ""
    sb = get_supabase()
    if not sb:
        return True, ""
    resp = sb.table("clientes").select("plano, ia_usado").eq("id", user_id).execute()
    if not resp.data:
        return True, ""
    dados = resp.data[0]
    if dados.get("plano") == "premium":
        return True, ""
    if dados.get("ia_usado"):
        return False, "Já usaste a tua análise gratuita com IA. Faz upgrade para o Premium para gerar propostas ilimitadas."
    return True, ""

# --- TELA DE AUTENTICAÇÃO ---
def tela_login():
    st.markdown("""
    <div style='text-align:center; margin-bottom:2rem; display:flex; justify-content:center;'>
        <svg width="140" height="140" viewBox="0 0 200 200" xmlns="http://www.w3.org/2000/svg">
            <path d="M65 62 L138 88 L102 100 L88 136 Z" fill="#2F81F7"/>
            <text x="100" y="178" font-family="Inter, sans-serif" font-size="30" font-weight="700" letter-spacing="3" fill="#E6EDF3" text-anchor="middle">MIRA</text>
        </svg>
    </div>
    """, unsafe_allow_html=True)
    
    col_centered = st.columns([1, 2, 1])[1] if not st.session_state.get("is_mobile", False) else [st]
    
    with col_centered:
        sb = get_supabase()

        if sb:
            try:
                site_url = st.secrets.get("SITE_URL", "")
            except Exception:
                site_url = ""
            if st.button("🔵 Continuar com Google", use_container_width=True, key="btn_google"):
                try:
                    resp_oauth = sb.auth.sign_in_with_oauth({
                        "provider": "google",
                        "options": {"redirect_to": site_url} if site_url else {}
                    })
                    st.markdown(f"<meta http-equiv='refresh' content='0; url={resp_oauth.url}'>", unsafe_allow_html=True)
                    st.link_button("Clica aqui se não fores redirecionado", resp_oauth.url, use_container_width=True)
                except Exception as e:
                    st.error(f"Erro ao iniciar login com Google: {e}")
            st.markdown("<p style='text-align:center; opacity:0.5; margin: 0.5rem 0;'>ou</p>", unsafe_allow_html=True)

        aba_entrar, aba_criar = st.tabs(["Entrar", "Criar conta"])

        with aba_entrar:
            email = st.text_input("E-mail", key="login_email")
            senha = st.text_input("Senha", type="password", key="login_senha")
            manter_sessao = st.checkbox("Manter sessão iniciada neste dispositivo", value=True, key="manter_sessao")
            if st.button("Entrar no Painel", type="primary", key="btn_entrar", use_container_width=True):
                if sb:
                    try:
                        resp = sb.auth.sign_in_with_password({"email": email, "password": senha})
                        garantir_registro_cliente(resp.user.id, email)
                        st.session_state["autenticado"] = True
                        st.session_state["chave_cliente"] = resp.user.id
                        st.session_state["email_cliente"] = email
                        dados_perfil = buscar_perfil_completo(resp.user.id)
                        st.session_state["nome_cliente"] = dados_perfil["apelido"] or email.split("@")[0]
                        st.session_state["perfil_apelido"] = dados_perfil["apelido"]
                        st.session_state["perfil_empresa"] = dados_perfil["empresa"]
                        st.session_state["perfil_oferta"] = dados_perfil["perfil_oferta"]
                        st.session_state["perfil_carregado"] = True

                        if manter_sessao and resp.session:
                            cookies["refresh_token"] = resp.session.refresh_token
                            cookies.save()

                        st.rerun()
                    except Exception:
                        st.error("E-mail ou senha incorretos.")
                else:
                    # Fallback local para desenvolvimento
                    st.session_state["autenticado"] = True
                    st.session_state["chave_cliente"] = "demo"
                    st.session_state["email_cliente"] = email or "demo@saas.com"
                    st.session_state["nome_cliente"] = email.split("@")[0] if email else "Utilizador"
                    st.rerun()

            with st.expander("Esqueci a senha"):
                email_recuperar = st.text_input("O teu e-mail", key="email_recuperar")
                if st.button("Enviar link de recuperação", key="btn_recuperar", use_container_width=True):
                    if sb and email_recuperar:
                        try:
                            sb.auth.reset_password_email(email_recuperar)
                            st.success("Se esse e-mail estiver registado, foi enviado um link de recuperação. Confere também o Spam.")
                        except Exception as e:
                            st.error(f"Erro ao enviar: {e}")
                    else:
                        st.warning("Escreve o teu e-mail primeiro.")

        with aba_criar:
            novo_email = st.text_input("E-mail", key="cad_email")
            novo_apelido = st.text_input("Nome para usar nas propostas", key="cad_apelido", placeholder="Ex: João Manuel")
            nova_empresa = st.text_input("Onde trabalhas / nome da tua empresa", key="cad_empresa", placeholder="Ex: JM Marketing Digital")
            nova_oferta = st.text_area("O que fazes / que serviço ofereces", key="cad_oferta",
                                        placeholder="Ex: Faço gestão de redes sociais e criação de sites.", height=80)
            nova_senha = st.text_input("Senha", type="password", key="cad_senha")
            if st.button("Criar Conta", type="primary", key="btn_criar", use_container_width=True):
                if sb:
                    try:
                        resp = sb.auth.sign_up({"email": novo_email, "password": nova_senha})
                        if resp.user:
                            garantir_registro_cliente(resp.user.id, novo_email, novo_apelido, nova_empresa, nova_oferta)
                        st.success("Conta criada! Confirma o e-mail enviado.")
                    except Exception as e:
                        st.error(f"Erro ao criar conta: {e}")
                else:
                    st.info("Configura as chaves do Supabase para guardar utilizadores reais.")

    st.stop()

def tela_definir_nova_senha(token_hash):
    st.markdown("<h2 style='text-align:center;'>Definir Nova Senha</h2>", unsafe_allow_html=True)
    col = st.columns([1, 2, 1])[1]
    with col:
        nova = st.text_input("Nova senha (mínimo 6 caracteres)", type="password", key="nova_senha_1")
        confirmar = st.text_input("Confirma a nova senha", type="password", key="nova_senha_2")
        if st.button("Guardar nova senha", type="primary", use_container_width=True):
            if len(nova) < 6:
                st.error("A senha precisa ter pelo menos 6 caracteres.")
            elif nova != confirmar:
                st.error("As senhas não coincidem.")
            else:
                sb = get_supabase()
                try:
                    resp = sb.auth.verify_otp({"token_hash": token_hash, "type": "recovery"})
                    if resp.session:
                        sb.auth.set_session(resp.session.access_token, resp.session.refresh_token)
                    sb.auth.update_user({"password": nova})

                    # Já entra direto no painel, aproveitando a sessão que acabou de ser validada
                    user_id = resp.user.id
                    email = resp.user.email
                    dados_perfil = buscar_perfil_completo(user_id)
                    st.session_state["autenticado"] = True
                    st.session_state["chave_cliente"] = user_id
                    st.session_state["email_cliente"] = email
                    st.session_state["nome_cliente"] = dados_perfil["apelido"] or email.split("@")[0]
                    st.session_state["perfil_apelido"] = dados_perfil["apelido"]
                    st.session_state["perfil_empresa"] = dados_perfil["empresa"]
                    st.session_state["perfil_oferta"] = dados_perfil["perfil_oferta"]
                    st.session_state["perfil_carregado"] = True

                    st.query_params.clear()
                    st.success("Senha atualizada! A entrar...")
                    st.rerun()
                except Exception as e:
                    st.error(f"Este link pode ter expirado ou já foi usado. Pede um novo link de recuperação. ({e})")
    st.stop()

def restaurar_sessao_do_cookie():
    """Tenta religar a sessão usando o refresh_token guardado no cookie do navegador."""
    refresh_token = cookies.get("refresh_token")
    if not refresh_token:
        return False
    sb = get_supabase()
    if not sb:
        return False
    try:
        resp = sb.auth.refresh_session(refresh_token)
        dados_perfil = buscar_perfil_completo(resp.user.id)
        st.session_state["autenticado"] = True
        st.session_state["chave_cliente"] = resp.user.id
        st.session_state["email_cliente"] = resp.user.email
        st.session_state["nome_cliente"] = dados_perfil["apelido"] or resp.user.email.split("@")[0]
        st.session_state["perfil_apelido"] = dados_perfil["apelido"]
        st.session_state["perfil_empresa"] = dados_perfil["empresa"]
        st.session_state["perfil_oferta"] = dados_perfil["perfil_oferta"]
        st.session_state["perfil_carregado"] = True
        # Atualiza o cookie com o novo refresh_token (Supabase rotaciona a cada uso)
        if resp.session:
            cookies["refresh_token"] = resp.session.refresh_token
            cookies.save()
        return True
    except Exception:
        # Cookie inválido/expirado - limpa e segue pro login normal
        cookies["refresh_token"] = ""
        cookies.save()
        return False

if "autenticado" not in st.session_state:
    _qp = st.query_params
    if _qp.get("type") == "recovery" and _qp.get("token_hash"):
        tela_definir_nova_senha(_qp.get("token_hash"))

    # Retorno do login com Google: pode vir como ?code=... (PKCE) ou access_token/refresh_token (implicit)
    if _qp.get("code"):
        sb = get_supabase()
        try:
            resp = sb.auth.exchange_code_for_session({"auth_code": _qp.get("code")})
            dados_perfil = buscar_perfil_completo(resp.user.id)
            garantir_registro_cliente(resp.user.id, resp.user.email, dados_perfil.get("apelido"))
            st.session_state["autenticado"] = True
            st.session_state["chave_cliente"] = resp.user.id
            st.session_state["email_cliente"] = resp.user.email
            st.session_state["nome_cliente"] = dados_perfil["apelido"] or resp.user.email.split("@")[0]
            st.session_state["perfil_apelido"] = dados_perfil["apelido"]
            st.session_state["perfil_empresa"] = dados_perfil["empresa"]
            st.session_state["perfil_oferta"] = dados_perfil["perfil_oferta"]
            st.session_state["perfil_carregado"] = True
            if resp.session:
                cookies["refresh_token"] = resp.session.refresh_token
                cookies.save()
            st.query_params.clear()
            st.rerun()
        except Exception as e:
            st.error(f"Erro ao entrar com Google: {e}")

    elif _qp.get("access_token") and _qp.get("refresh_token"):
        sb = get_supabase()
        try:
            resp = sb.auth.set_session(_qp.get("access_token"), _qp.get("refresh_token"))
            dados_perfil = buscar_perfil_completo(resp.user.id)
            garantir_registro_cliente(resp.user.id, resp.user.email, dados_perfil.get("apelido"))
            st.session_state["autenticado"] = True
            st.session_state["chave_cliente"] = resp.user.id
            st.session_state["email_cliente"] = resp.user.email
            st.session_state["nome_cliente"] = dados_perfil["apelido"] or resp.user.email.split("@")[0]
            st.session_state["perfil_apelido"] = dados_perfil["apelido"]
            st.session_state["perfil_empresa"] = dados_perfil["empresa"]
            st.session_state["perfil_oferta"] = dados_perfil["perfil_oferta"]
            st.session_state["perfil_carregado"] = True
            if resp.session:
                cookies["refresh_token"] = resp.session.refresh_token
                cookies.save()
            st.query_params.clear()
            st.rerun()
        except Exception as e:
            st.error(f"Erro ao entrar com Google: {e}")

    if not restaurar_sessao_do_cookie():
        tela_login()

# --- FUNÇÕES DE NEGÓCIO E RASTREIO ---
def limpar_para_pdf(texto):
    return str(texto).replace("&", "e").replace("<", "").replace(">", "") if texto else "N/A"

def extrair_email_do_site(url):
    if not url or url == "N/A": return ""
    try:
        headers = {'User-Agent': 'Mozilla/5.0'}
        resposta = requests.get(url, headers=headers, timeout=3)
        emails = set(re.findall(r"[a-zA-Z0-9._%+-]+@[a-zA-Z0-9.-]+\.[a-zA-Z]{2,}", resposta.text))
        validos = [e for e in emails if not e.endswith(('.png', '.jpg', '.jpeg', '.gif', '.webp'))]
        return ", ".join(validos[:2])
    except Exception:
        return ""

def gerar_link_whatsapp(telefone, mensagem=""):
    if not telefone: return ""
    num = re.sub(r'\D', '', str(telefone))
    if not num: return ""
    from urllib.parse import quote
    return f"https://wa.me/{num}?text={quote(mensagem)}" if mensagem else f"https://wa.me/{num}"

def extrair_mensagem_da_analise(analise_ia):
    if not analise_ia: return ""
    match = re.search(r"MENSAGEM:?\s*(.+?)(?:\n\n_\(modelo|\Z)", analise_ia, re.DOTALL | re.IGNORECASE)
    return match.group(1).strip() if match else analise_ia.split("_(modelo")[0].strip()

def enviar_email_gmail(remetente, senha_app, destinatario, assunto, corpo):
    import smtplib
    from email.mime.text import MIMEText
    if not remetente or not senha_app:
        return False, "Insere as tuas credenciais do Gmail na barra lateral."
    if not destinatario:
        return False, "Nenhum e-mail de destino encontrado."

    msg = MIMEText(corpo)
    msg["Subject"] = assunto
    msg["From"] = remetente
    msg["To"] = destinatario

    try:
        with smtplib.SMTP_SSL("smtp.gmail.com", 465, timeout=15) as servidor:
            servidor.login(remetente, senha_app)
            servidor.sendmail(remetente, [destinatario], msg.as_string())
        return True, "E-mail enviado com sucesso!"
    except Exception as e:
        return False, f"Erro ao enviar: {e}"

def calcular_score_oportunidade(site, avaliacao, num_avaliacoes, telefone):
    score = 0
    if not site: score += 30
    try:
        nota = float(avaliacao) if avaliacao else 0.0
        if 0 < nota < 4.0: score += 20
    except Exception: pass
    if not telefone: score -= 10
    return score

def temperatura_lead(score):
    if score >= 40: return "🔴 Quente"
    if score >= 15: return "🟡 Morno"
    return "🔵 Frio"

# --- STATUS DO LEAD (mini-CRM) ---
def carregar_status_leads(user_id, nomes):
    sb = get_supabase()
    if not sb or not nomes: return {}
    resp = sb.table("leads_status").select("nome_empresa, status").eq("user_id", user_id).in_("nome_empresa", nomes).execute()
    return {r["nome_empresa"]: r["status"] for r in resp.data} if resp.data else {}

def salvar_status_lead(user_id, nome_empresa, status):
    sb = get_supabase()
    if not sb: return
    sb.table("leads_status").upsert({
        "user_id": user_id, "nome_empresa": nome_empresa, "status": status,
        "atualizado_em": "now()"
    }, on_conflict="user_id,nome_empresa").execute()

def contar_status_mes(user_id):
    sb = get_supabase()
    if not sb: return {}
    resp = sb.table("leads_status").select("status").eq("user_id", user_id).execute()
    contagem = {"Novo": 0, "Contactado": 0, "Respondeu": 0, "Fechado": 0}
    for r in (resp.data or []):
        s = r.get("status", "Novo")
        contagem[s] = contagem.get(s, 0) + 1
    return contagem

def listar_leads_status(user_id):
    """Devolve todos os leads guardados do cliente, mais recentes primeiro."""
    sb = get_supabase()
    if not sb: return []
    resp = sb.table("leads_status").select("*").eq("user_id", user_id).order("atualizado_em", desc=True).execute()
    return resp.data or []

def dias_desde(data_iso):
    from datetime import datetime, timezone
    try:
        dt = datetime.fromisoformat(data_iso.replace("Z", "+00:00"))
        return (datetime.now(timezone.utc) - dt).days
    except Exception:
        return None

def _resposta_valida(texto):
    """Detecta se a IA devolveu um veredito de moderação em vez de texto de verdade."""
    if not texto or len(texto.strip()) < 15:
        return False
    marcas_moderacao = ["user safety", "response safety", "safety categories", '"rating"']
    texto_lower = texto.lower()
    return not any(marca in texto_lower for marca in marcas_moderacao)

def dica_ia_para_lead(nome_empresa, status, dias, perfil_oferta, api_key):
    if not api_key:
        return "Erro: Chave API OpenRouter necessária."
    headers = {"Authorization": f"Bearer {api_key}", "Content-Type": "application/json"}
    prompt = f"""
    Estou a prospectar a empresa '{nome_empresa}'. Status atual: {status}.
    Já se passaram {dias if dias is not None else '?'} dias desde a última atualização.
    Eu ofereço: {perfil_oferta or 'serviços diversos'}.
    Dá uma sugestão curta e prática (máximo 3 frases) do que fazer a seguir com este lead,
    considerando o status e o tempo parado. Responde em Português, direto ao ponto.
    """
    for modelo in ["openrouter/free", "google/gemma-4-31b-it:free", "nvidia/nemotron-3-super-120b-a12b:free"]:
        try:
            payload = {"model": modelo, "messages": [{"role": "user", "content": prompt}]}
            res = requests.post(OPENROUTER_API_URL, headers=headers, json=payload, timeout=25)
            if res.status_code == 200:
                texto = res.json()['choices'][0]['message']['content'].strip()
                if _resposta_valida(texto):
                    return texto
        except Exception:
            continue
    return "Não foi possível gerar a dica agora."

# --- INTEGRAÇÃO IA ---
def analisar_com_ia(nome, nicho, site, avaliacao, objetivo, api_key, remetente_nome="Um consultor"):
    if not api_key: return "Erro: Chave API OpenRouter necessária.", "Chave API em falta"

    headers = {
        "Authorization": f"Bearer {api_key}",
        "Content-Type": "application/json",
        "HTTP-Referer": "https://streamlit.io",
        "X-Title": "Motor B2B"
    }

    prompt = f"""
    Objetivo: '{objetivo}'.
    Empresa: {nome} (Nicho: {nicho}). Website: {'Sim' if site else 'Não'}. Avaliação: {avaliacao}.
    Quem está a enviar a mensagem se chama: {remetente_nome}.
    Retorna APENAS:
    1. DIAGNÓSTICO: (1 frase curta)
    2. MENSAGEM: (1 mensagem persuasiva de WhatsApp para abordagem, assinada com o nome "{remetente_nome}" no lugar de qualquer placeholder tipo "[Seu Nome]")
    """

    erros_debug = []
    for modelo in ["openrouter/free", "google/gemma-4-31b-it:free", "nvidia/nemotron-3-super-120b-a12b:free"]:
        try:
            payload = {"model": modelo, "messages": [{"role": "user", "content": prompt}]}
            res = requests.post(OPENROUTER_API_URL, headers=headers, json=payload, timeout=25)
            if res.status_code == 200:
                texto = res.json()['choices'][0]['message']['content'].strip()
                if _resposta_valida(texto):
                    return texto, None
                else:
                    erros_debug.append(f"{modelo}: resposta inválida ({texto[:60]})")
            else:
                erros_debug.append(f"{modelo}: HTTP {res.status_code} - {res.text[:150]}")
        except Exception as e:
            erros_debug.append(f"{modelo}: exceção - {e}")

    return "Não foi possível gerar análise no momento.", " | ".join(erros_debug)

# --- BUSCA EM CASCATA: OSM (grátis) -> RapidAPI (barato) -> Google (caro, último recurso) ---
def buscar_lugares_osm(nicho, regiao, limit, progress=None):
    try:
        geo = requests.get("https://nominatim.openstreetmap.org/search", 
                           params={"q": regiao, "format": "json", "limit": 1}, 
                           headers={"User-Agent": "scraper-app/1.0"}, timeout=15).json()
        if not geo: return []
        lat, lon = float(geo[0]["lat"]), float(geo[0]["lon"])

        query = f"""
        [out:json][timeout:25];
        (
          node["name"~"{nicho}",i](around:20000,{lat},{lon});
          node["shop"](around:20000,{lat},{lon})["name"~"{nicho}",i];
        );
        out body {limit * 2};
        """
        resp = requests.post("https://overpass-api.de/api/interpreter", data={"data": query}, timeout=30)
        if resp.status_code != 200:
            st.info(f"OpenStreetMap indisponível (erro {resp.status_code}). Seguindo com outras fontes.")
            return []

        resultados = []
        vistos = set()
        for el in resp.json().get("elements", []):
            if len(resultados) >= limit: break
            tags = el.get("tags", {})
            nome = tags.get("name")
            if not nome or nome in vistos: continue
            vistos.add(nome)

            tel = tags.get("phone", tags.get("contact:phone", ""))
            site = tags.get("website", tags.get("contact:website", ""))
            score = calcular_score_oportunidade(site, "", 0, tel)

            resultados.append({
                "Score": score, "Nome": nome, "Telefone": tel,
                "E-mail": extrair_email_do_site(site) if site else "",
                "Site": site, "Avaliação": "N/A", "Fonte": "OpenStreetMap"
            })
            if progress: progress(len(resultados), limit, "A buscar (OpenStreetMap, grátis)...")
        return resultados
    except Exception:
        return []

def buscar_lugares(query, api_key, limit, nicho, regiao, progress=None):
    """Motor RapidAPI - pago (~$1,25-2,50/1000), usado só pro que faltar depois do OSM."""
    headers = {"x-rapidapi-key": api_key, "x-rapidapi-host": "local-business-data.p.rapidapi.com"}
    params = {"query": query, "limit": str(limit), "language": "pt"}
    try:
        res = requests.get(SEARCH_URL, headers=headers, params=params, timeout=20)
        if res.status_code != 200:
            if res.status_code == 429:
                st.warning("RapidAPI: limite mensal estourado (erro 429). Seguindo só com fontes grátis.")
            elif res.status_code in (401, 403):
                st.warning(f"RapidAPI: chave inválida ou sem permissão (erro {res.status_code}).")
            else:
                st.warning(f"RapidAPI retornou erro {res.status_code}.")
            return []

        dados = res.json().get("data", [])
        resultados = []
        for lugar in dados:
            if len(resultados) >= limit: break
            nome = lugar.get("name", "N/A")
            tel = lugar.get("phone_number", "")
            site = lugar.get("website", "")
            aval = lugar.get("rating", "")
            score = calcular_score_oportunidade(site, aval, lugar.get("review_count", 0), tel)
            resultados.append({
                "Score": score, "Nome": nome, "Telefone": tel,
                "E-mail": extrair_email_do_site(site), "Site": site,
                "Avaliação": str(aval), "Fonte": "RapidAPI"
            })
            if progress: progress(len(resultados), limit, "A extrair dados (RapidAPI)...")
        return resultados
    except Exception:
        return []

def buscar_lugares_google(nicho, regiao, limit, google_api_key, progress=None):
    """Motor Google Places (New) - o mais caro (~$32-40/1000), só entra se as outras 2 não bastarem."""
    if not google_api_key: return []
    try:
        url = "https://places.googleapis.com/v1/places:searchText"
        headers = {
            "Content-Type": "application/json",
            "X-Goog-Api-Key": google_api_key,
            "X-Goog-FieldMask": "places.displayName,places.nationalPhoneNumber,places.websiteUri,places.rating"
        }
        body = {"textQuery": f"{nicho} em {regiao}", "maxResultCount": min(limit, 20)}
        res = requests.post(url, headers=headers, json=body, timeout=15)
        if res.status_code != 200: return []

        resultados = []
        for lugar in res.json().get("places", [])[:limit]:
            nome = lugar.get("displayName", {}).get("text", "N/A")
            tel = lugar.get("nationalPhoneNumber", "")
            site = lugar.get("websiteUri", "")
            aval = lugar.get("rating", "")
            score = calcular_score_oportunidade(site, aval, 0, tel)
            resultados.append({
                "Score": score, "Nome": nome, "Telefone": tel,
                "E-mail": extrair_email_do_site(site) if site else "",
                "Site": site, "Avaliação": str(aval), "Fonte": "Google Places"
            })
            if progress: progress(len(resultados), limit, "A buscar (Google Places, pago)...")
        return resultados
    except Exception:
        return []

def buscar_leads_cascata(nicho, regiao, limit, rapidapi_key, google_api_key="", progress=None):
    """Ordem: OSM (grátis) -> RapidAPI (barato) -> Google (caro). Cada camada só busca o que falta."""
    resultados = buscar_lugares_osm(nicho, regiao, limit, progress)
    vistos = {r["Nome"].lower() for r in resultados}

    faltam = limit - len(resultados)
    if faltam > 0 and rapidapi_key:
        extras = buscar_lugares(f"{nicho} em {regiao}", rapidapi_key, faltam, nicho, regiao, progress)
        for r in extras:
            if r["Nome"].lower() not in vistos:
                resultados.append(r); vistos.add(r["Nome"].lower())

    faltam = limit - len(resultados)
    if faltam > 0 and google_api_key:
        extras = buscar_lugares_google(nicho, regiao, faltam, google_api_key, progress)
        for r in extras:
            if r["Nome"].lower() not in vistos:
                resultados.append(r); vistos.add(r["Nome"].lower())

    return sorted(resultados, key=lambda x: x["Score"], reverse=True)

# --- EXPORTAÇÃO: WORD E PDF ---
def criar_word(dados, nicho, regiao):
    doc = Document()
    doc.add_heading(f"Relatório de Prospeção: {nicho} em {regiao}", 0)
    for item in dados:
        doc.add_heading(item["Nome"], level=2)
        doc.add_paragraph(f"• Contato: {item['Telefone']} | Email: {item['E-mail']}")
        doc.add_paragraph(f"• Score Comercial: {item['Score']}")
        doc.add_paragraph(f"• Análise IA:\n{item.get('Análise IA', 'Não analisado.')}")
        doc.add_paragraph("-" * 30)
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer

def criar_pdf(dados, nicho, regiao):
    buffer = BytesIO()
    doc = SimpleDocTemplate(buffer, pagesize=letter)
    styles = getSampleStyleSheet()
    elementos = [Paragraph(f"Relatório Estratégico: {nicho} em {regiao}", styles['Title']), Spacer(1, 15)]
    for item in dados:
        nome = limpar_para_pdf(item['Nome'])
        elementos.append(Paragraph(f"<b>{nome}</b> (Score: {item['Score']})", styles['Heading2']))
        elementos.append(Paragraph(f"<b>Telefone:</b> {limpar_para_pdf(item['Telefone'])}", styles['Normal']))
        elementos.append(Paragraph(f"<b>E-mail:</b> {limpar_para_pdf(item['E-mail'])}", styles['Normal']))
        analise = limpar_para_pdf(item.get('Análise IA', 'Não gerado.'))
        elementos.append(Paragraph(f"<b>Análise IA:</b> {analise}", styles['Normal']))
        elementos.append(Spacer(1, 10))
    doc.build(elementos)
    buffer.seek(0)
    return buffer

# --- INTERFACE PRINCIPAL ---
c_title, c_user = st.columns([3, 1])
with c_title:
    st.markdown("""
    <svg width="90" height="90" viewBox="0 0 200 200" xmlns="http://www.w3.org/2000/svg">
        <path d="M65 62 L138 88 L102 100 L88 136 Z" fill="#2F81F7"/>
        <text x="100" y="178" font-family="Inter, sans-serif" font-size="30" font-weight="700" letter-spacing="3" fill="#E6EDF3" text-anchor="middle">MIRA</text>
    </svg>
    """, unsafe_allow_html=True)
    st.caption("Encontra os clientes certos, sem perder tempo.")

with c_user:
    if st.button(f"👤 {st.session_state.get('nome_cliente', 'Utilizador')}", key="btn_abrir_perfil", use_container_width=True):
        st.session_state["vista"] = "perfil"
        st.rerun()
    if st.button("📋 Meus Leads", key="btn_abrir_leads", use_container_width=True):
        st.session_state["vista"] = "leads"
        st.rerun()
    if st.button("Sair", key="btn_logout"):
        cookies["refresh_token"] = ""
        cookies.save()
        st.session_state.clear()
        st.rerun()

# --- PÁGINA DE PERFIL ---
if st.session_state.get("vista") == "perfil":
    st.markdown("## Perfil Profissional")
    st.caption("Estes dados são usados pela IA para personalizar as tuas propostas automaticamente.")

    if "perfil_carregado" not in st.session_state:
        dados_perfil = buscar_perfil_completo(st.session_state.get("chave_cliente", ""))
        st.session_state["perfil_apelido"] = dados_perfil["apelido"]
        st.session_state["perfil_empresa"] = dados_perfil["empresa"]
        st.session_state["perfil_oferta"] = dados_perfil["perfil_oferta"]
        st.session_state["perfil_carregado"] = True

    novo_apelido = st.text_input("Nome para usar nas propostas", value=st.session_state["perfil_apelido"],
                                  placeholder="Ex: João Manuel")
    nova_empresa = st.text_input("Onde trabalhas / nome da tua empresa", value=st.session_state["perfil_empresa"],
                                  placeholder="Ex: JM Marketing Digital")
    nova_oferta = st.text_area("O que fazes / que serviço ofereces", value=st.session_state["perfil_oferta"],
                                placeholder="Ex: Faço gestão de redes sociais e criação de sites para pequenos negócios.",
                                height=100)

    col_salvar, col_voltar = st.columns(2)
    with col_salvar:
        if st.button("Guardar perfil", type="primary", use_container_width=True):
            salvar_perfil_completo(st.session_state.get("chave_cliente", ""), novo_apelido, nova_empresa, nova_oferta)
            st.session_state["perfil_apelido"] = novo_apelido
            st.session_state["perfil_empresa"] = nova_empresa
            st.session_state["perfil_oferta"] = nova_oferta
            st.session_state["nome_cliente"] = novo_apelido or st.session_state.get("nome_cliente")
            st.success("Perfil guardado!")
    with col_voltar:
        if st.button("← Voltar à busca", use_container_width=True):
            st.session_state["vista"] = "busca"
            st.rerun()

    st.stop()  # não mostra o resto da página enquanto estiver na vista de perfil

# --- PÁGINA MEUS LEADS (histórico + assistente IA) ---
if st.session_state.get("vista") == "leads":
    st.markdown("## 📋 Meus Leads")
    st.caption("Histórico de todas as empresas que já marcaste com um status — com dicas da IA sobre o que fazer a seguir.")

    if st.button("← Voltar à busca", key="voltar_de_leads"):
        st.session_state["vista"] = "busca"
        st.rerun()

    todos_leads = listar_leads_status(st.session_state.get("chave_cliente", ""))

    if not todos_leads:
        st.info("Ainda não marcaste nenhum lead. Faz uma busca, gera propostas com IA, e muda o status de alguma empresa para começar a aparecer aqui.")
    else:
        contagem = contar_status_mes(st.session_state.get("chave_cliente", ""))
        k1, k2, k3, k4 = st.columns(4)
        k1.metric("Novo", contagem.get("Novo", 0))
        k2.metric("Contactado", contagem.get("Contactado", 0))
        k3.metric("Respondeu", contagem.get("Respondeu", 0))
        k4.metric("Fechado", contagem.get("Fechado", 0))

        st.divider()

        filtro = st.multiselect("Filtrar por status", ["Novo", "Contactado", "Respondeu", "Fechado"],
                                 default=["Novo", "Contactado", "Respondeu", "Fechado"])

        for lead in todos_leads:
            if lead["status"] not in filtro:
                continue
            dias = dias_desde(lead.get("atualizado_em", ""))
            texto_dias = f"há {dias} dia(s)" if dias is not None else ""

            with st.expander(f"{lead['nome_empresa']} — {lead['status']} ({texto_dias})"):
                col_status, col_dica = st.columns(2)

                with col_status:
                    novo_status = st.selectbox(
                        "Atualizar status", ["Novo", "Contactado", "Respondeu", "Fechado"],
                        index=["Novo", "Contactado", "Respondeu", "Fechado"].index(lead["status"]),
                        key=f"status_hist_{lead['id']}"
                    )
                    if novo_status != lead["status"]:
                        salvar_status_lead(st.session_state.get("chave_cliente", ""), lead["nome_empresa"], novo_status)
                        st.rerun()

                with col_dica:
                    if st.button("💡 Pedir dica da IA", key=f"dica_{lead['id']}", use_container_width=True):
                        with st.spinner("A pensar no próximo passo..."):
                            dica = dica_ia_para_lead(
                                lead["nome_empresa"], lead["status"], dias,
                                st.session_state.get("perfil_oferta", ""), OPENROUTER_API_KEY
                            )
                        st.info(dica)

    st.stop()

# Sidebar de Configurações do Gmail
with st.sidebar:
    st.subheader("⚙️ Configurações de Envio")
    GMAIL_ENDERECO = st.text_input("O teu Gmail", placeholder="exemplo@gmail.com")
    GMAIL_APP_PASSWORD = st.text_input("Senha de App do Gmail", type="password", help="Gera nas configurações de Segurança da Google (Senhas de App).")
    st.caption("Credenciais mantidas de forma temporária nesta sessão.")

st.markdown("<hr style='border-color: rgba(255,255,255,0.08); margin: 1rem 0;'>", unsafe_allow_html=True)

# Busca o plano do cliente uma vez, pra ajustar os limites da interface
if "plano_cliente" not in st.session_state:
    st.session_state["plano_cliente"] = buscar_plano_cliente(st.session_state.get("chave_cliente", ""))
eh_premium = st.session_state["plano_cliente"] == "premium" or st.session_state.get("email_cliente", "") in EMAILS_ADMIN

if not eh_premium:
    st.caption("🆓 Plano Grátis: 3 buscas/dia · até 6 empresas por busca · 1 uso da IA. [Upgrade para Premium →](#)")

# Formulário de Busca
col1, col2, col3 = st.columns([2, 2, 1])
nicho = col1.text_input("Nicho / Setor", placeholder="Ex: Clínicas, Restaurantes")
regiao = col2.text_input("Região", placeholder="Ex: Maputo, Lisboa")
limite_qtd = 50 if eh_premium else MAX_LEADS_FREE
max_leads = col3.number_input("Qtd. Máxima", min_value=1, max_value=limite_qtd, value=min(10, limite_qtd))

if st.button("Iniciar Varredura 🚀", type="primary", use_container_width=True):
    if not nicho or not regiao:
        st.warning("Preenche o Nicho e a Região para pesquisar.")
    else:
        permitido, msg_uso = verificar_e_registrar_uso(st.session_state.get("chave_cliente", "demo"), st.session_state.get("email_cliente", ""))
        if not permitido:
            st.error(msg_uso)
        else:
            # Reforço no servidor: mesmo que alguém force o campo, o free nunca passa de MAX_LEADS_FREE
            max_leads_real = max_leads if eh_premium else min(max_leads, MAX_LEADS_FREE)
            bar = st.progress(0, "Iniciando...")
            resultados = buscar_leads_cascata(nicho, regiao, max_leads_real, RAPIDAPI_KEY, GOOGLE_PLACES_API_KEY,
                                              lambda c, t, m: bar.progress(min(c/t, 1.0), m))
            bar.empty()
            if resultados:
                st.session_state.update({'leads': resultados, 'n': nicho, 'r': regiao})
            else:
                st.info("Nenhuma empresa encontrada para essa região com este termo.")

# Exibição de Resultados
if 'leads' in st.session_state:
    df = pd.DataFrame(st.session_state['leads'])
    df["Temperatura"] = df["Score"].apply(temperatura_lead)

    # Painel de KPIs
    kpi1, kpi2, kpi3, kpi4 = st.columns(4)
    kpi1.metric("Leads Encontrados", len(df))
    kpi2.metric("Quentes 🔴", (df["Temperatura"] == "🔴 Quente").sum())
    kpi3.metric("Com E-mail", (df["E-mail"] != "").sum())

    contagem_status = contar_status_mes(st.session_state.get("chave_cliente", ""))
    kpi4.metric("Fechados (total)", contagem_status.get("Fechado", 0))

    st.markdown(f"### 🎯 Oportunidades Encontradas ({len(df)})")
    st.dataframe(df, use_container_width=True)

    if "status_leads" not in st.session_state:
        nomes = [l["Nome"] for l in st.session_state['leads']]
        st.session_state["status_leads"] = carregar_status_leads(st.session_state.get("chave_cliente", ""), nomes)

    st.markdown("---")
    st.markdown("### 🧠 Gerar Abordagens com IA")

    tem_perfil = bool(st.session_state.get("perfil_oferta", "").strip())
    objetivo_digitado = st.text_area(
        "Objetivo específico para esta busca (opcional)",
        value="",
        placeholder=f"Deixa em branco para usar o teu perfil: \"{st.session_state.get('perfil_oferta', '')[:70]}...\"" if tem_perfil else "Ex: Vender gestão de redes sociais",
        height=70
    )
    objetivo_atual = objetivo_digitado.strip() or st.session_state.get("perfil_oferta", "").strip()
    if not objetivo_digitado.strip() and tem_perfil:
        st.caption("✓ Usando o teu perfil guardado (nenhum objetivo específico digitado acima).")

    if len(df) <= 1:
        qtd = len(df)
        st.caption(f"Vai analisar {qtd} empresa encontrada.")
    else:
        qtd = st.slider("Quantidade de empresas para analisar:", 1, len(df), min(3, len(df)))

    if st.button("Gerar Propostas com IA 🤖", type="primary"):
        pode_ia, msg_ia = verificar_ia_disponivel(st.session_state.get("chave_cliente", ""), st.session_state.get("email_cliente", ""))
        if not pode_ia:
            st.error(msg_ia)
        else:
            st.session_state["debug_ia"] = []
            with st.spinner("A IA está a redigir as mensagens de abordagem..."):
                for i in range(qtd):
                    empresa = st.session_state['leads'][i]
                    resp, debug = analisar_com_ia(empresa["Nome"], st.session_state['n'], empresa["Site"], 
                                         empresa["Avaliação"], objetivo_atual, OPENROUTER_API_KEY,
                                         remetente_nome=st.session_state.get('nome_cliente', 'Um consultor'))
                    st.session_state['leads'][i]["Análise IA"] = resp
                    if debug:
                        st.session_state["debug_ia"].append(f"{empresa['Nome']}: {debug}")
                    elif not eh_premium:
                        marcar_ia_usada(st.session_state.get("chave_cliente", ""))
                st.rerun()

    if st.session_state.get("debug_ia"):
        with st.expander("⚠️ Detalhes do erro da IA (DEBUG)", expanded=True):
            for linha in st.session_state["debug_ia"]:
                st.code(linha)
            if st.button("Fechar aviso de debug"):
                st.session_state["debug_ia"] = []
                st.rerun()

    # Cards de Envio (Glassmorphism)
    leads_com_analise = [l for l in st.session_state['leads'] if l.get("Análise IA")]
    if leads_com_analise:
        st.markdown("### 📩 Enviar Abordagens")
        for idx, empresa in enumerate(st.session_state['leads']):
            if not empresa.get("Análise IA"): continue
            
            mensagem = extrair_mensagem_da_analise(empresa["Análise IA"])
            with st.expander(f"📍 {empresa['Nome']} (Score: {empresa['Score']} · {temperatura_lead(empresa['Score'])})"):
                st.markdown(empresa["Análise IA"])
                msg_editada = st.text_area("Mensagem de Envio", value=mensagem, key=f"msg_{idx}", height=100)

                status_atual = st.session_state.get("status_leads", {}).get(empresa["Nome"], "Novo")
                novo_status = st.selectbox("Status", ["Novo", "Contactado", "Respondeu", "Fechado"],
                                            index=["Novo", "Contactado", "Respondeu", "Fechado"].index(status_atual),
                                            key=f"status_{idx}")
                if novo_status != status_atual:
                    salvar_status_lead(st.session_state.get("chave_cliente", ""), empresa["Nome"], novo_status)
                    st.session_state.setdefault("status_leads", {})[empresa["Nome"]] = novo_status
                    st.rerun()

                colA, colB = st.columns(2)
                
                # E-mail Directo
                with colA:
                    if empresa.get("E-mail"):
                        if st.button(f"Enviar por E-mail 📧", key=f"email_{idx}", use_container_width=True):
                            ok, status = enviar_email_gmail(GMAIL_ENDERECO, GMAIL_APP_PASSWORD, 
                                                            empresa["E-mail"], f"Proposta para {empresa['Nome']}", msg_editada)
                            st.success(status) if ok else st.error(status)
                    else:
                        st.caption("Sem e-mail detetado.")

                # WhatsApp Directo
                with colB:
                    link_wa = gerar_link_whatsapp(empresa.get("Telefone", ""), msg_editada)
                    if link_wa:
                        st.link_button("Abrir WhatsApp 💬", link_wa, use_container_width=True)
                    else:
                        st.caption("Sem telefone válido.")

    st.markdown("---")
    st.markdown("### 📤 Exportar Relatórios")
    df_final = pd.DataFrame(st.session_state['leads'])
    b1, b2, b3 = st.columns(3)
    b1.download_button("↓ Excel (CSV)", data=df_final.to_csv(index=False).encode("utf-8"),
                       file_name="leads.csv", mime="text/csv", use_container_width=True)
    b2.download_button("↓ Documento Word", data=criar_word(st.session_state['leads'], st.session_state['n'], st.session_state['r']),
                       file_name="leads.docx", mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document", use_container_width=True)
    b3.download_button("↓ Relatório PDF", data=criar_pdf(st.session_state['leads'], st.session_state['n'], st.session_state['r']),
                       file_name="leads.pdf", mime="application/pdf", use_container_width=True)
